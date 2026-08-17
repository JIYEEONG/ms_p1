from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(r"C:\Users\user\Desktop\ms_p1\deliverables\ZERO_dashboard_improvement_and_contribution_report.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = "18202A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "697582"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
ORANGE = "FF4500"
SOFT_ORANGE = "FFF0EA"
WHITE = "FFFFFF"
BORDER = "DCE3EA"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(0.492)
sec.footer_distance = Inches(0.492)

def set_font(run, name="Calibri", size=11, color=NAVY, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(NAVY)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for style_name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
]:
    st = styles[style_name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for name in ["List Bullet", "List Number"]:
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    st.font.size = Pt(11)
    st.paragraph_format.left_indent = Inches(0.5)
    st.paragraph_format.first_line_indent = Inches(-0.25)
    st.paragraph_format.space_after = Pt(8)
    st.paragraph_format.line_spacing = 1.167

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths_dxa)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcW = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[idx]))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)

def format_cell(cell, bold=False, color=NAVY, size=9.5):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05
        for r in p.runs:
            set_font(r, size=size, color=color, bold=bold)

def add_table(headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        shade(table.rows[0].cells[i], PALE_BLUE)
        format_cell(table.rows[0].cells[i], bold=True, color=DARK_BLUE, size=9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            format_cell(cells[i], size=9.2)
    set_table_geometry(table, widths)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return table

def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.75)
    r = p.add_run(text)
    set_font(r)
    return p

def add_number(text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_font(r)
    return p

def add_callout(label, text, fill=SOFT_ORANGE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + "  ")
    set_font(r, size=10.5, color=ORANGE, bold=True)
    r = p.add_run(text)
    set_font(r, size=10.5, color=NAVY)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

# Running header/footer
header = sec.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_font(hp.add_run("ZERO Dashboard | Improvement & Contribution Report"), size=9, color=MUTED, bold=True)
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(fp.add_run("2026.08.17  |  Internal Project Record"), size=8.5, color=MUTED)

# Cover / memo masthead
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(6)
r = p.add_run("PROJECT IMPROVEMENT REPORT")
set_font(r, size=10, color=ORANGE, bold=True)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
r = p.add_run("ZERO 판매·재고 의사결정 대시보드\n개선 내역 및 사용자 기여 보고서")
set_font(r, size=24, color=NAVY, bold=True)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(20)
r = p.add_run("5dt-team2-Next.js-project-2nd와 ms_p1의 연속 개발 과정을 기준으로 정리")
set_font(r, size=12.5, color=MUTED)

meta = [
    ("작성 기준일", "2026년 8월 17일"),
    ("대상 시스템", "ZERO 판매·재고 의사결정 대시보드"),
    ("검토 범위", "대화 기반 요구사항, 현재 코드 구조, Git 이력 및 작업 트리"),
    ("문서 목적", "개선 성과와 사용자의 제품·UX 기여를 재사용 가능한 기록으로 정리"),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(label + ": "), size=10.5, color=NAVY, bold=True)
    set_font(p.add_run(value), size=10.5, color=MUTED)

add_callout("핵심 결론", "단순 모니터링 화면이었던 대시보드를 매출 수익성, 목표 달성, 재고 위험, HUB 이동, 판매 예측과 발주 판단까지 연결하는 실무형 의사결정 도구로 확장했다. 사용자는 반복적인 시각 검토를 통해 정보 위계와 업무 흐름의 품질 기준을 직접 정의했다.")

doc.add_heading("1. 프로젝트 식별 및 검토 기준", level=1)
doc.add_paragraph("두 폴더는 동일 제품의 연속 개발본으로 확인된다. 원격 저장소 명칭은 다르지만 현재 두 작업 폴더 모두 같은 최신 커밋을 가리키며, 기존 프로젝트에는 ms_p1 저장소가 personal 원격으로 연결되어 있다.")
add_table(
    ["구분", "5dt-team2-Next.js-project-2nd", "ms_p1"],
    [
        ("현재 브랜치", "main", "main"),
        ("확인 커밋", "4cbdfd1", "4cbdfd1"),
        ("원격 관계", "기존 팀 저장소 + personal(ms_p1)", "JIYEEONG/ms_p1"),
        ("판단", "동일 코드베이스의 원 프로젝트", "개인 연속 개발·개선본"),
    ],
    [1800, 3780, 3780],
)
doc.add_paragraph("보고서의 개선 내역은 (1) 대화에서 확정된 요구사항, (2) 현재 수정 파일과 컴포넌트 구조, (3) Git 커밋 메시지를 교차 확인해 작성했다. 일부 순이익·비용 값은 사용자의 허용에 따라 가상 산정 로직을 포함한다.")

doc.add_heading("2. 경영진 요약", level=1)
add_table(
    ["개선 축", "이전 문제", "개선 결과"],
    [
        ("매출 판단", "매출 중심, 긴 숫자와 분산된 KPI", "순이익·비용·단가·전년 비교까지 한 화면에서 판단"),
        ("재고 운영", "요약 숫자와 상세 업무 흐름의 연결 부족", "HUB/상품 전용 페이지, 위험·WOS·이동·담당 처리까지 연결"),
        ("예측/발주", "예측 차트와 결과 해석이 분리", "선택 상품·기간·예상 판매·부족 시점·발주 결과를 통합"),
        ("필터 사용성", "스크롤 후 조건을 다시 수정하기 어려움", "상단 공통 필터 + 원위치 편집 가능한 스크롤 요약"),
        ("시각 체계", "페이지별 반경·표면·선택색 불일치", "공통 radius, surface hierarchy, 브랜드 오렌지 시스템 확립"),
        ("기술 안정성", "포트·Docker DB·API 오류 처리 불일치", "8000 포트 통일, 로컬 SQL 인증서 대응, 오류 상태 구분"),
    ],
    [1700, 3300, 4360],
)

doc.add_heading("3. 개선 과정: 요구 발견에서 시스템화까지", level=1)
for item in [
    "정보 추가: 매출·순이익·총비용·개당 손익, 재고 위험, 예측 결과 등 업무 판단에 필요한 데이터를 보강했다.",
    "정보 구조 재정렬: 중요한 값부터 읽히도록 KPI 순서, 표 열 순서, 제목-필터-본문 순서를 반복 조정했다.",
    "상호작용 개선: 카테고리 클릭 이동, HUB 카드 선택, 필터 즉시 편집, 정렬·기간 토글 등 다음 행동으로 이어지게 했다.",
    "시각 언어 통합: 글래스 재질 실험을 거쳐 최종적으로 White/Cool Gray/Pale Blue 90%, Soft Orange 7%, #FF4500 3% 체계를 확정했다.",
    "회귀 수정: 툴팁 잘림, 사이드바·필터 겹침, 표 헤더 모서리, 화면 밖 팝오버, ESLint 상태 업데이트 문제를 지속 점검했다.",
]: add_bullet(item)

doc.add_heading("4. 매출 현황 페이지 개선", level=1)
doc.add_heading("4.1 목표와 달성", level=2)
for item in [
    "핵심 순서를 목표 달성률, 매출, 순이익으로 재배치하고 주문 건수·판매량·ASP·ATV·UPT를 보조 KPI로 분리했다.",
    "ASP(평균 판매 단가), ATV(주문당 평균 결제금액), UPT(주문당 평균 판매수량)를 영문+한국어로 병기하고 계산식을 유지했다.",
    "요약 단위(억원·만원·만건) 아래에 정확한 전체 숫자를 병기해 빠른 인지와 검증 가능성을 동시에 확보했다.",
    "최종 컬러 가이드에 따라 목표 달성률은 중립적으로 유지하고, 의사결정상 중요한 순이익만 #FF4500으로 강조했다.",
]: add_bullet(item)

doc.add_heading("4.2 순매출 추이", level=2)
for item in [
    "해당 기간·전년 동기·재작년 동기를 2행 라벨과 서로 다른 선 스타일로 구분했다.",
    "툴팁 안에서 동일한 1·2·3단계 텍스트 위계를 적용하고, 데이터 없음은 0원이 아닌 대시(—)로 표현했다.",
    "전년 대비 증감률과 목표 대비 값을 비교 지점에 배치하고, 화면 가장자리에서 툴팁이 잘리지 않도록 여유와 위치를 조정했다.",
    "현재 기간은 브랜드 오렌지, 비교 기간은 Pale Blue/Cool Gray로 바꿔 데이터 의미와 색상 의미를 연결했다.",
]: add_bullet(item)

doc.add_heading("4.3 카테고리별 매출 Top 5", level=2)
for item in [
    "Top 10에서 Top 5로 압축하고 대분류·중분류·소분류 선택 및 매출 순·순이익 순 정렬을 제공했다.",
    "열을 순위/카테고리, 순이익, 매출, 총비용, 개당 순이익, 개당 비용 순으로 구성했다.",
    "원가·인건비·물류비·마케팅비·관리비 등 상품 한 개에 투입되는 전체 비용을 비용 상세로 확인할 수 있게 했다.",
    "긴 숫자는 요약 단위와 정확한 금액을 함께 표시하고 모든 열을 왼쪽 기준의 시각적 마진으로 정렬했다.",
    "카테고리명을 클릭하면 상품별 재고 페이지로 이동하도록 연결해 분석에서 실행으로 이어지는 흐름을 만들었다.",
    "표 헤더와 첫 행의 잘림을 방지하고 헤더 양쪽 상단 모서리를 공통 radius로 마감했다.",
]: add_bullet(item)

doc.add_heading("5. 필터 및 탐색 경험 개선", level=1)
for item in [
    "모든 주요 페이지의 카테고리·검색·기간 필터를 페이지 최상단에 배치했다.",
    "매출 현황에는 기간·대/중분류·시즌·HUB, HUB 재고에는 HUB·검색·분류·정렬, 상품 재고에는 검색·위험·분류·SKU, 예측에는 분류·상품을 구성했다.",
    "원래 필터가 화면에 보일 때는 요약을 숨기고, 스크롤로 원래 필터가 사라질 때만 작은 요약 탭이 나타나도록 했다.",
    "요약의 수정 버튼을 눌러도 페이지 상단으로 이동하지 않고 현재 위치에서 필터를 바로 변경하도록 했다.",
    "사이드바와 상단 필터가 겹치지 않도록 위치와 가로 폭을 조정하고, 선택 메뉴와 활성 탭에 동일한 선택색 규칙을 적용했다.",
]: add_bullet(item)

doc.add_heading("6. HUB별 재고 관리 개선", level=1)
for item in [
    "허브별 재고 요약과 상세 영역을 하나의 업무 페이지로 정리하고, 제목·설명을 필터보다 위에 배치했다.",
    "보유재고, 가용재고, 예약재고, 이동 중 수량을 숫자로 통일해 요약 카드 간 비교가 가능해졌다.",
    "각 HUB 카드를 누르면 해당 HUB의 보유 상품 목록으로 이동하고 선택 상태를 명확히 표시한다.",
    "HUB별 상품 목록에서 대/중/소 카테고리와 재고 많은 순·적은 순 정렬을 지원한다.",
    "안전재고, 부족·초과 수량, 일평균 판매량, 예상 품절일, 입고 예정 수량·날짜, 타 HUB 재고, 최근 7일 증감률, 담당자·처리 상태를 추가했다.",
    "수도권 통합 허브 보유 상품 다음에 HUB별 소진 속도 및 재고 상태가 이어지도록 업무 읽기 순서를 조정했다.",
]: add_bullet(item)

doc.add_heading("7. 상품별 재고 관리 개선", level=1)
for item in [
    "상품별 재고 요약과 상품·SKU 목록의 위험 상태 명칭을 동일하게 맞춰 같은 지표가 다른 이름으로 보이는 문제를 줄였다.",
    "가용재고, 품절 임박 SKU, 과잉재고 SKU, 장기재고 SKU를 요약하고, 요약 카드를 눌러 해당 상태만 필터링할 수 있게 했다.",
    "상품·SKU·옵션별 판매 속도, 안전재고, WOS, 판매율, 클레임을 확인하는 구조를 유지하면서 상세 업무 정보를 확장했다.",
    "페이지 제목과 설명을 카테고리 필터 위로 이동해 제목→조건→결과의 일관된 읽기 순서를 만들었다.",
]: add_bullet(item)

doc.add_heading("8. 판매 예측 및 발주 제안 개선", level=1)
for item in [
    "현재 선택 상품을 박스 형태로 먼저 보여줘 무엇에 대한 예측인지 즉시 알 수 있게 했다.",
    "예측 기간을 판매 예측 및 발주 제안 영역 안으로 이동하고 1~4주 선택 결과를 별도 조작 없이 하단에 표시했다.",
    "상품별 판매 예측 TOP 10에 가장 잘 팔릴 상품·가장 안 팔릴 상품·재고 부족 예상 모드를 통합했다.",
    "기존의 개별 예상 판매량·권장 발주량 카드를 정리하고, 상품·현재+입고·주간 예상판매·부족 시점·위험도를 행 단위 결과로 재구성했다.",
    "예상 기말재고는 그래프 대신 현재 가용재고 + 입고 예정 - 예상 판매량 = 예상 기말재고의 숫자 중심 구조로 변경했다.",
    "미연결 상태였던 날씨 보정 문구와 불필요한 계산식 안내를 제거해 실제 제공 기능과 화면 설명을 일치시켰다.",
    "실제 판매는 Pale Blue, 예측은 #FF4500으로 표시해 차트와 결과표의 의미색을 일관되게 했다.",
]: add_bullet(item)

doc.add_heading("9. 디자인 시스템과 접근성 개선", level=1)
add_table(
    ["원칙", "적용 내용"],
    [
        ("표면 위계", "Page < Outer Container < Card < Input/Select 순으로 명도 차이를 설계"),
        ("반경", "페이지별 사각/과도한 pill 혼용을 줄이고 공통 radius 체계를 적용"),
        ("선 사용", "명도 차이가 부족할 때만 1px 저대비 border를 보조적으로 사용"),
        ("텍스트 위계", "제목·핵심 값·보조 설명의 3단계 위계를 페이지 위치와 무관하게 통일"),
        ("정렬", "표와 카드의 텍스트를 시각적 마진 기준으로 정렬하고 동일 gutter를 적용"),
        ("선택 상태", "사이드바·탭·필터의 활성 상태를 Soft Orange 배경과 Orange indicator로 통일"),
        ("색상", "Neutral 90% + Soft Orange 7% + #FF4500 3%, 의미색은 경고/성공/오류와 분리"),
        ("상태 전달", "상승·하락은 색상뿐 아니라 ▲/▼ 기호를 함께 사용"),
    ],
    [2000, 7360],
)

doc.add_heading("10. 백엔드·데이터 연결 및 품질 개선", level=1)
for item in [
    "프론트엔드 /backend rewrite가 127.0.0.1:8000 백엔드로 연결되는 구조를 점검하고 팀 기준 포트를 8000으로 통일했다.",
    "로컬 Docker SQL Server의 인증서 환경을 지원하도록 연결 설정을 보완했다.",
    "기간·대/중분류·시즌·HUB 필터가 KPI, 추이, 카테고리 집계에 동일하게 전달되도록 API 파라미터를 정리했다.",
    "네트워크 오류와 API 오류를 구분하고 로딩·빈 상태·데이터 없음 상태를 구체적으로 표시했다.",
    "React effect 안의 동기 setState 경고를 피하도록 사용자 이벤트 시점에 로딩 상태를 시작하고 비동기 완료 시 종료하는 구조로 개선했다.",
    "주요 변경마다 Next.js production build와 TypeScript 검사를 수행해 기능 회귀를 확인했다.",
]: add_bullet(item)

doc.add_heading("11. 사용자의 역량 분석", level=1)
add_callout("종합 평가", "사용자는 화면을 예쁘게 만드는 수준을 넘어, 현업 의사결정의 순서와 데이터 의미를 UI 구조로 번역하는 제품 설계 역량을 보여줬다.", fill=PALE_BLUE)
add_table(
    ["역량", "관찰된 행동", "프로젝트 가치"],
    [
        ("제품 사고", "‘그래서 무엇을 봐야 하는가’를 중심으로 KPI와 화면 순서를 재설계", "모니터링 화면을 행동 가능한 도구로 전환"),
        ("정보 구조 설계", "제목→필터→결과, 목표율→매출→순이익 등 읽기 순서를 반복 조정", "인지 부하와 탐색 시간을 감소"),
        ("데이터 문해력", "매출·순이익·비용·개당 손익·WOS·ASP/ATV/UPT를 연결", "매출 규모와 수익성·재고 위험을 함께 판단"),
        ("도메인 이해", "안전재고, 품절일, 입고 예정, 타 HUB 이동, 담당 상태를 요구", "패션/리테일 재고 운영의 실제 업무 반영"),
        ("UX 디테일", "툴팁 잘림, 시각적 마진, gutter, 탭 위치, 스크롤 편집을 구체적으로 검수", "사용 중 발생하는 마찰을 세밀하게 제거"),
        ("디자인 시스템", "명도 우선, border 보조, surface level, radius와 선택색을 전역 규칙으로 정의", "페이지 간 일관성과 확장성 확보"),
        ("접근성 감수성", "색상 단독 의존을 피하고 기호·라벨·명도 차이를 함께 사용", "다양한 사용자에게 상태 전달 강화"),
        ("반복 개선", "레퍼런스를 검토하고 필요 시 원복하며 최종 컬러 가이드를 확정", "실험을 통제된 학습 과정으로 전환"),
    ],
    [1700, 4000, 3660],
)

doc.add_heading("12. 사용자의 핵심 기여 지점", level=1)
for title, body in [
    ("1) 문제 정의", "정보량이 충분해도 판단 속도가 느릴 수 있다는 본질적 문제를 제기하고, ‘무엇을 먼저 봐야 하는가’를 개선의 기준으로 설정했다."),
    ("2) 지표 및 업무 요구 정의", "순이익·전체 비용·개당 손익, 안전재고·품절일·타 HUB 재고, 예측·권장 발주 등 필요한 데이터 항목을 구체화했다."),
    ("3) 정보 우선순위 결정", "Top 10을 Top 5로 압축하고, 표 열·KPI·페이지 제목과 필터의 순서를 직접 결정했다."),
    ("4) 상호작용 설계", "스크롤 필터 요약의 노출 조건, 현재 위치 편집, 카테고리→상품 재고 이동, HUB 선택 후 상세 이동을 정의했다."),
    ("5) 시각 품질 기준 수립", "명도·여백·모서리·재질·색상 비율을 문서 수준의 규칙으로 발전시켜 디자인 시스템의 방향을 제공했다."),
    ("6) 품질 검수 및 원복 판단", "겹침·잘림·낮은 대비·과한 효과를 스크린샷으로 특정하고, 의도와 다른 시안은 원복하도록 결정했다."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    set_font(p.add_run(title + " — "), bold=True, color=DARK_BLUE)
    set_font(p.add_run(body), color=NAVY)

doc.add_heading("13. 기여의 성격: 사용자가 사실상 수행한 역할", level=1)
add_table(
    ["역할", "수행한 기여"],
    [
        ("Product Owner", "업무 목표, 우선순위, 삭제/유지 범위, 최종 승인 기준 결정"),
        ("UX Designer", "정보 구조, 인터랙션, 상태 노출, 탐색 경로, 오류·빈 상태 정의"),
        ("UI/Design System Reviewer", "컬러 비율, 표면 단계, radius, 텍스트 위계, 선택 상태 규칙 수립"),
        ("Domain Expert", "매출 수익성, 상품 비용, 재고 위험, HUB 이동, 발주 판단에 필요한 지표 제공"),
        ("QA Lead", "실제 화면 캡처로 잘림·겹침·대비·정렬 문제를 재현하고 회귀 여부 확인"),
    ],
    [2500, 6860],
)

doc.add_heading("14. 대표 성과와 기대 효과", level=1)
for item in [
    "의사결정 속도: 매출·순이익·비용·재고 위험을 한 흐름에서 확인해 별도 계산과 화면 이동을 줄인다.",
    "업무 실행력: 분석 결과에서 상품 상세, HUB 재고, 발주 제안으로 바로 이동할 수 있다.",
    "데이터 신뢰: 요약 단위와 정확한 값, 계산식과 비교 기준을 함께 제공한다.",
    "운영 효율: 필터를 다시 찾지 않고 스크롤 위치에서 즉시 수정할 수 있다.",
    "확장성: 공통 surface, border, radius, selected state 규칙으로 새 페이지도 같은 언어로 확장할 수 있다.",
]: add_bullet(item)

doc.add_heading("15. 현재 상태와 후속 권고", level=1)
doc.add_paragraph("현재 작업 트리에는 프론트엔드·백엔드 다수 파일의 미커밋 변경이 존재한다. 기능과 스타일의 범위가 넓으므로 다음 단계에서는 변경을 목적별로 분리해 안정적으로 인계하는 것이 좋다.")
for item in [
    "기능, API, 스타일 변경을 각각 별도 커밋으로 정리하고 스크린샷 기반 회귀 기준을 남긴다.",
    "가상 순이익·비용 산정 항목의 정의와 계산 기준을 데이터 사전으로 명문화한다.",
    "필터 조합별 API 응답, 빈 상태, 모바일 폭, 툴팁 경계 조건을 자동 테스트로 추가한다.",
    "실제 사용자 업무 시나리오(매출 부진 탐색, 품절 대응, HUB 이동, 발주 결정)로 사용성 테스트를 진행한다.",
    "브랜드 오렌지 3% 원칙과 의미색 분리를 디자인 토큰 문서로 고정한다.",
]: add_number(item)

doc.add_heading("부록 A. 주요 변경 파일", level=1)
add_table(
    ["영역", "대표 파일"],
    [
        ("전역 UI 시스템", "src/app/globals.css, src/app/page.tsx"),
        ("매출 현황", "OverviewTab.tsx, KpiCard.tsx, SalesTrendChart.tsx, CategorySalesChart.tsx"),
        ("필터", "FilterBar.tsx, FilterSummaryBar.tsx"),
        ("재고", "HubInventory.tsx, ProductInventory.tsx, InventoryStatus.tsx"),
        ("예측", "ForecastInventory.tsx"),
        ("탐색/레이아웃", "Header.tsx, Sidebar.tsx, ScrollControls.tsx"),
        ("API/백엔드", "dashboardApi.ts, backend/app/api/v1/dashboard.py, schemas/dashboard.py, core/config.py"),
    ],
    [2400, 6960],
)

doc.add_heading("부록 B. 검증 메모", level=1)
doc.add_paragraph("보고서 작성 시 확인한 최신 Git 커밋은 4cbdfd1(fix: support local Docker SQL Server certificate)이며, 이전 이력에는 재고 대시보드 개선, 포트 8000 통일, 필터 연동과 API 오류 처리, 역할별 접근 제어, AI 리포트 개선 등이 포함되어 있다. 본 보고서는 현재 코드 상태와 대화에서 합의된 변경을 중심으로 작성했으며, 정량 성과는 별도 사용자 테스트 또는 운영 로그가 확보되면 추가 측정해야 한다.")

doc.core_properties.title = "ZERO 대시보드 개선 내역 및 사용자 기여 보고서"
doc.core_properties.subject = "5dt-team2-Next.js-project-2nd / ms_p1 연속 개선 기록"
doc.core_properties.author = "OpenAI Codex"
doc.core_properties.keywords = "ZERO, dashboard, UX, inventory, sales, contribution"
doc.save(OUT)
print(OUT)
